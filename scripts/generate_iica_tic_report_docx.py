from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
LOGO_PATH = ROOT / "public" / "logos" / "official" / "iica.png"
OUTPUT_PATH = ROOT / "docs" / "Informe-Gobernanza-Gasto-TIC-IICA-Chile-2026.docx"


def clp(value: int) -> str:
    return f"CLP {value:,}".replace(",", ".")


def add_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def make_header_row(table, labels: list[str]) -> None:
    row = table.rows[0]
    for idx, label in enumerate(labels):
        cell = row.cells[idx]
        cell.text = label
        add_cell_shading(cell, "1F4E78")
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(10)


def add_table_style(table) -> None:
    table.style = "Table Grid"
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)


def add_title_page(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(0.9)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    if LOGO_PATH.exists():
        p_logo = doc.add_paragraph()
        p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_logo.add_run().add_picture(str(LOGO_PATH), width=Inches(2.0))

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("Informe Técnico-Financiero\nGobernanza y Gasto TIC")
    run_title.bold = True
    run_title.font.size = Pt(24)
    run_title.font.color.rgb = RGBColor(15, 76, 129)

    p_subtitle = doc.add_paragraph()
    p_subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_subtitle = p_subtitle.add_run("Representación IICA Chile")
    run_subtitle.bold = True
    run_subtitle.font.size = Pt(16)
    run_subtitle.font.color.rgb = RGBColor(42, 42, 42)

    doc.add_paragraph("")
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_run = info.add_run(
        "Horizonte presupuestario: 12 meses\n"
        "Moneda: Peso chileno (CLP)\n"
        "Audiencia objetivo: Dirección y Finanzas"
    )
    info_run.font.size = Pt(11)

    doc.add_paragraph("")
    p_date = doc.add_paragraph()
    p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = p_date.add_run(f"Santiago de Chile, {date.today().strftime('%d-%m-%Y')}")
    date_run.italic = True
    date_run.font.size = Pt(11)

    doc.add_page_break()


def add_h1(doc: Document, text: str) -> None:
    heading = doc.add_heading(text, level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in heading.runs:
        run.font.color.rgb = RGBColor(15, 76, 129)


def add_h2(doc: Document, text: str) -> None:
    heading = doc.add_heading(text, level=2)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(30, 97, 149)


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text, style="List Bullet")
    for run in p.runs:
        run.font.size = Pt(11)


def add_body(doc: Document) -> None:
    normal_style = doc.styles["Normal"]
    normal_style.font.name = "Calibri"
    normal_style.font.size = Pt(11)

    add_h1(doc, "1. Resumen ejecutivo")
    doc.add_paragraph(
        "Este informe propone una infraestructura de gobernanza de datos y de control del gasto TIC para la "
        "Representación IICA Chile, orientada a sostener la operación de plataformas digitales, reducir riesgo "
        "operacional y habilitar decisiones presupuestarias basadas en evidencia."
    )
    add_bullet(doc, "Decisión 1: aprobar presupuesto TIC anual centralizado con control mensual y umbrales de escalamiento.")
    add_bullet(doc, "Decisión 2: crear un Departamento de Gobernanza y Transformacion Digital (DGTD) con mandato formal sobre datos, IA, arquitectura y control de cartera.")
    add_bullet(doc, "Decisión 3: implementar práctica FinOps con reporting mensual en CLP y revisión de desvíos.")
    add_bullet(doc, "Monto referencial anual recomendado (escenario base operativo): CLP 12.660.000.")

    add_h1(doc, "2. Objetivo, alcance y criterios de éxito")
    doc.add_paragraph(
        "El alcance considera un modelo por capas: (i) núcleo de plataformas y datos de proyectos digitales en curso, "
        "(ii) operación y gobierno TIC, y (iii) capa transversal TIC opcional para servicios institucionales comunes."
    )
    add_bullet(doc, "Horizonte: 12 meses.")
    add_bullet(doc, "Moneda de gobierno: CLP.")
    add_bullet(doc, "Audiencia principal: Dirección y Finanzas.")
    add_bullet(doc, "Criterio de éxito: aprobación de presupuesto, comité de gobernanza y mecanismo de control operativo mensual.")

    add_h1(doc, "3. Estado operacional actual (as-is)")
    doc.add_paragraph(
        "La plataforma opera sobre una arquitectura web con Next.js, base de datos y servicios gestionados en Supabase, "
        "consumo de IA con Gemini y ejecución de automatizaciones mediante GitHub Actions. El repositorio evidencia una "
        "cadencia operativa diaria y semanal para ingesta y descubrimiento de información."
    )

    table_ops = doc.add_table(rows=1, cols=4)
    add_table_style(table_ops)
    make_header_row(table_ops, ["Proceso", "Frecuencia", "Objetivo", "Evidencia interna"])
    rows_ops = [
        (
            "Ingesta de scrapers",
            "Diaria (03:00 Chile)",
            "Actualizar base de oportunidades y marcar obsolescencia",
            ".github/workflows/ingest-scrapers.yml",
        ),
        (
            "AI Discovery",
            "Semanal (lunes)",
            "Descubrimiento asistido con guardrails de validación",
            ".github/workflows/discover-projects.yml",
        ),
        (
            "Backfill embeddings",
            "Operación periódica",
            "Completar vectorización para búsqueda semántica",
            "scripts/backfill-embeddings.ts",
        ),
    ]
    for row_data in rows_ops:
        row = table_ops.add_row().cells
        for idx, value in enumerate(row_data):
            row[idx].text = value

    add_h1(doc, "4. Modelo de gobernanza de datos y TIC")
    doc.add_paragraph(
        "Se recomienda evolucionar desde un esquema basado solo en comites a una capacidad institucional permanente: "
        "un Departamento de Gobernanza y Transformacion Digital (DGTD) con atribuciones formales, comites de apoyo "
        "y un sistema de control de costo-riesgo-valor. El diseño se alinea con ISO/IEC 38500:2024 y practicas FinOps."
    )

    add_h2(doc, "4.1 Modelo institucional propuesto: DGTD")
    doc.add_paragraph(
        "El DGTD se propone como unidad liviana pero con responsabilidad end-to-end sobre portafolio digital, "
        "gobernanza de datos, control FinOps y gestion de riesgo tecnologico. Los comites dejan de ser la estructura "
        "principal y pasan a ser instancias de validacion y escalamiento."
    )

    table_roles = doc.add_table(rows=1, cols=5)
    add_table_style(table_roles)
    make_header_row(table_roles, ["Nivel", "Estructura", "Mandato principal", "Dotacion referencial", "Dependencia"])
    rows_roles = [
        (
            "Estrategico",
            "Jefatura DGTD",
            "Define estrategia digital anual, metas de costo/valor y criterios de priorizacion institucional",
            "1 jefatura",
            "Direccion de Representacion",
        ),
        (
            "Tactico",
            "Unidad FinOps y Portafolio",
            "Consolida presupuesto TIC, opera showback, forecast trimestral y oficina de proyectos digitales (PMO)",
            "2 analistas",
            "DGTD",
        ),
        (
            "Tactico",
            "Unidad Arquitectura, Datos e IA Responsable",
            "Administra politicas de datos/IA, controles de acceso, trazabilidad, continuidad y estandares de arquitectura",
            "2 especialistas",
            "DGTD",
        ),
        (
            "Operacion",
            "Enlaces TIC por area",
            "Canalizan demanda, validan beneficios, levantan riesgos y aseguran adopcion de practicas",
            "0.2 FTE por area",
            "Area funcional + DGTD",
        ),
    ]
    for row_data in rows_roles:
        row = table_roles.add_row().cells
        for idx, value in enumerate(row_data):
            row[idx].text = value

    add_h2(doc, "4.2 Comites de apoyo (segunda linea de decision)")
    table_committees = doc.add_table(rows=1, cols=4)
    add_table_style(table_committees)
    make_header_row(table_committees, ["Instancia", "Proposito", "Cadencia", "Resultado esperado"])
    committee_rows = [
        (
            "Consejo Directivo TIC-Finanzas",
            "Aprobar decisiones de alto impacto (presupuesto, priorizacion, riesgos criticos)",
            "Mensual",
            "Acta con decisiones estrategicas y responsables",
        ),
        (
            "PMO Digital Operativa",
            "Controlar avance, hitos, gasto real y medidas de correccion",
            "Quincenal",
            "Backlog priorizado y plan de acciones",
        ),
        (
            "Mesa de Riesgo de Datos e IA",
            "Evaluar casos de uso IA, cumplimiento y excepciones de datos",
            "Mensual",
            "Registro de riesgos, excepciones y mitigaciones",
        ),
    ]
    for row_data in committee_rows:
        row = table_committees.add_row().cells
        for idx, value in enumerate(row_data):
            row[idx].text = value

    add_h2(doc, "4.3 Modelo de decisiones y RACI clave")
    table_raci = doc.add_table(rows=1, cols=5)
    add_table_style(table_raci)
    make_header_row(table_raci, ["Decision", "A (Accountable)", "R (Responsible)", "C (Consulted)", "I (Informed)"])
    raci_rows = [
        (
            "Presupuesto TIC anual y reforecast trimestral",
            "Direccion",
            "Jefatura DGTD",
            "Finanzas, PMO Digital",
            "Areas usuarias",
        ),
        (
            "Priorizacion de iniciativas digitales",
            "Jefatura DGTD",
            "Unidad FinOps y Portafolio",
            "Lideres de area",
            "Consejo Directivo",
        ),
        (
            "Excepcion por sobreconsumo >20%",
            "Direccion",
            "Jefatura DGTD",
            "Finanzas",
            "Consejo Directivo",
        ),
        (
            "Aprobacion de uso de datos sensibles en IA",
            "Jefatura DGTD",
            "Unidad Arquitectura, Datos e IA",
            "Legal/Compliance",
            "Direccion",
        ),
        (
            "Plan de continuidad y pruebas de recuperacion",
            "Jefatura DGTD",
            "Unidad Arquitectura, Datos e IA",
            "TI operativa",
            "Direccion y Finanzas",
        ),
    ]
    for row_data in raci_rows:
        row = table_raci.add_row().cells
        for idx, value in enumerate(row_data):
            row[idx].text = value

    add_h2(doc, "4.4 Politicas minimas")
    add_bullet(doc, "Clasificación de datos por criticidad y reglas de tratamiento por nivel.")
    add_bullet(doc, "Control de acceso, trazabilidad y revisión periódica de privilegios.")
    add_bullet(doc, "Política de uso de IA: datos permitidos/no permitidos, retención y registro de prompts críticos.")
    add_bullet(doc, "Plan de continuidad: respaldo, recuperación y pruebas periódicas de restauración.")
    add_bullet(doc, "Politica de arquitectura objetivo con estandares de integracion, observabilidad y seguridad por defecto.")

    add_h1(doc, "5. Mecanismo presupuestario: Fondo TIC central")
    doc.add_paragraph(
        "Se propone un Fondo TIC central con reglas de asignación transparentes, seguimiento mensual y capacidad de "
        "corrección ante variaciones de demanda o precio."
    )
    add_bullet(doc, "Distribución recomendada del fondo: 65% Run (operación), 20% Grow (mejora), 15% Reserva de riesgo.")
    add_bullet(doc, "Modelo de imputación: costos comunes centralizados + showback mensual por proyecto/unidad.")
    add_bullet(doc, "Escalamiento financiero: desvío >10% revisión directiva; desvío >20% plan de contención inmediato.")
    add_bullet(doc, "Administracion del fondo: DGTD (ejecucion) + Finanzas (control) + Direccion (aprobacion de excepciones).")

    add_h1(doc, "6. Costeo técnico por capas y por operación")
    doc.add_paragraph(
        "Los montos son referenciales y deben calibrarse con facturación real histórica. El objetivo es proveer un "
        "marco defendible para aprobación y seguimiento."
    )

    table_costs = doc.add_table(rows=1, cols=5)
    add_table_style(table_costs)
    make_header_row(
        table_costs,
        ["Capa", "Componente", "Costo mensual", "Costo anual", "Observación"],
    )
    cost_rows = [
        ("Núcleo plataforma", "Vercel (hosting/app)", clp(210000), clp(2520000), "Incluye crecimiento moderado"),
        ("Núcleo plataforma", "Supabase (DB/Auth/Storage)", clp(260000), clp(3120000), "Prod + staging"),
        ("Núcleo plataforma", "Gemini (uso IA)", clp(90000), clp(1080000), "Uso controlado por cuotas"),
        ("Núcleo plataforma", "GitHub Actions / CI", clp(70000), clp(840000), "Pipelines + artefactos"),
        ("Operación y gobierno", "Observabilidad y soporte", clp(190000), clp(2280000), "Monitoreo y soporte operativo"),
        ("Operación y gobierno", "Seguridad y cumplimiento", clp(140000), clp(1680000), "Controles y hardening"),
        ("Operación y gobierno", "Reserva técnica", clp(95000), clp(1140000), "Incidentes y variaciones"),
        ("Transversal opcional", "Servicios TIC compartidos", clp(420000), clp(5040000), "Licencias/integraciones/soporte"),
    ]
    for row_data in cost_rows:
        row = table_costs.add_row().cells
        for idx, value in enumerate(row_data):
            row[idx].text = value

    add_h2(doc, "6.1 Escenarios presupuestarios (12 meses)")
    table_scenarios = doc.add_table(rows=1, cols=5)
    add_table_style(table_scenarios)
    make_header_row(
        table_scenarios,
        ["Escenario", "Descripción", "Costo mensual", "Costo anual", "Uso recomendado"],
    )
    scenarios = [
        (
            "Base operativo",
            "Capas núcleo + operación y gobierno",
            clp(1055000),
            clp(12660000),
            "Aprobación mínima inmediata",
        ),
        (
            "Integral",
            "Base + 60% capa transversal",
            clp(1307000),
            clp(15684000),
            "Operación institucional ampliada",
        ),
        (
            "Crecimiento",
            "Base + capa transversal completa + mayor reserva",
            clp(1875000),
            clp(22500000),
            "Escalamiento y nuevos proyectos",
        ),
    ]
    for row_data in scenarios:
        row = table_scenarios.add_row().cells
        for idx, value in enumerate(row_data):
            row[idx].text = value

    add_h1(doc, "7. Control mensual, KPIs y FinOps")
    add_bullet(doc, "KPI 1: Gasto real mensual vs presupuesto aprobado (CLP y % desvío).")
    add_bullet(doc, "KPI 2: Forecast de cierre trimestral de gasto TIC.")
    add_bullet(doc, "KPI 3: Costo unitario por operación (ingesta, discovery, búsqueda, soporte).")
    add_bullet(doc, "KPI 4: Consumo IA por caso de uso y costo por resultado útil.")
    add_bullet(doc, "KPI 5: Incidentes por sobreconsumo y tiempo medio de contención.")
    add_bullet(doc, "KPI 6: Porcentaje de costos asignados con trazabilidad por proyecto.")

    add_h2(doc, "7.1 Cadencia de control")
    table_cadence = doc.add_table(rows=1, cols=4)
    add_table_style(table_cadence)
    make_header_row(table_cadence, ["Hito", "Fecha objetivo", "Responsable", "Salida esperada"])
    cadence_rows = [
        ("Cierre TIC mensual", "T+3 hábil", "Finanzas + TI", "Reporte consolidado de gasto real"),
        ("Comité Operativo FinOps", "T+5 hábil", "TI + Finanzas", "Plan de acciones correctivas"),
        ("Comité Directivo", "Mensual", "Dirección", "Decisiones de ajuste/priorización"),
    ]
    for row_data in cadence_rows:
        row = table_cadence.add_row().cells
        for idx, value in enumerate(row_data):
            row[idx].text = value

    add_h1(doc, "8. Riesgos críticos y mitigaciones")
    table_risks = doc.add_table(rows=1, cols=4)
    add_table_style(table_risks)
    make_header_row(table_risks, ["Riesgo", "Impacto", "Mitigación", "Owner"])
    risks = [
        (
            "Sobrecosto por consumo inesperado",
            "Alto",
            "Alertas y límites de gasto por plataforma; playbook de contención en 24h",
            "Unidad FinOps y Portafolio (DGTD)",
        ),
        (
            "Uso inadecuado de datos en IA",
            "Alto",
            "Política de datos para IA + controles de anonimización + revisión mensual",
            "Unidad Arquitectura, Datos e IA (DGTD)",
        ),
        (
            "Opacidad de costos por proyecto",
            "Medio",
            "Showback mensual obligatorio y asignación por tags/centro de costo",
            "DGTD + Finanzas",
        ),
        (
            "Dependencia de proveedor",
            "Medio",
            "Plan de portabilidad de datos y revisión semestral de arquitectura",
            "Jefatura DGTD",
        ),
    ]
    for row_data in risks:
        row = table_risks.add_row().cells
        for idx, value in enumerate(row_data):
            row[idx].text = value

    add_h1(doc, "9. Hoja de ruta de implementación")
    add_h2(doc, "9.1 Primeros 90 días")
    add_bullet(doc, "Dia 1-30: emitir resolucion interna para crear DGTD, definir mandato, RACI y baseline de costos.")
    add_bullet(doc, "Dia 31-60: designar jefatura interina DGTD, activar PMO digital y tablero FinOps con alertas por umbrales.")
    add_bullet(doc, "Dia 61-90: ejecutar primer ciclo completo de control, reforecast y comite directivo con decisiones formalizadas.")
    add_h2(doc, "9.2 Ciclo anual")
    add_bullet(doc, "Reforecast trimestral.")
    add_bullet(doc, "Revisión semestral de riesgos y arquitectura.")
    add_bullet(doc, "Cierre anual con evaluación de valor generado y eficiencia del gasto TIC.")
    add_bullet(doc, "Evaluacion anual de madurez DGTD y ajuste de dotacion/capacidades.")

    add_h1(doc, "10. Solicitud de decisión")
    doc.add_paragraph(
        "Se solicita a Direccion y Finanzas aprobar: (i) la creacion del Departamento de Gobernanza y Transformacion "
        "Digital (DGTD), (ii) la implementacion del Fondo TIC central en escenario Base Operativo, y (iii) la transicion "
        "controlada al escenario Integral segun cumplimiento de KPIs y madurez de gestion de costos."
    )

    add_h1(doc, "11. Referencias técnicas y normativas")
    references = [
        "ISO/IEC 38500:2024 - https://www.iso.org/search.html?q=ISO%2FIEC%2038500%3A2024",
        "FinOps Framework - https://www.finops.org/framework/",
        "FOCUS Specification - https://focus.finops.org/focus-specification/",
        "NIST AI RMF - https://www.nist.gov/itl/ai-risk-management-framework",
        "NIST AI RMF 1.0 PDF - https://nvlpubs.nist.gov/nistpubs/ai/NIST.AI.100-1.pdf",
        "NIST AI 600-1 (GenAI profile) - https://nvlpubs.nist.gov/nistpubs/ai/NIST.AI.600-1.pdf",
        "OECD Recommendation - https://legalinstruments.oecd.org/en/instruments/OECD-LEGAL-0406",
        "Vercel Pricing - https://vercel.com/pricing",
        "Supabase Pricing - https://supabase.com/pricing",
        "Gemini API Pricing - https://ai.google.dev/gemini-api/docs/pricing",
        "GitHub Actions Billing - https://docs.github.com/en/billing/managing-billing-for-github-actions/about-billing-for-github-actions",
        "Tipo de cambio referencial - https://mindicador.cl/api",
    ]
    for item in references:
        add_bullet(doc, item)

    add_h2(doc, "Evidencia de operación en repositorio")
    evidence = [
        ".github/workflows/ingest-scrapers.yml",
        ".github/workflows/discover-projects.yml",
        "scripts/discover-projects.ts",
        "scripts/run-scrapers.ts",
        "scripts/backfill-embeddings.ts",
        "prisma/schema.prisma",
        "lib/utils/env.ts",
    ]
    for item in evidence:
        add_bullet(doc, item)


def main() -> None:
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    add_title_page(doc)
    add_body(doc)
    doc.save(OUTPUT_PATH)
    print(f"Documento generado en: {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
